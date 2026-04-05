import os, fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

receipts_dir = '/Users/luna-openclaw/Desktop/spesenabrechnungen/email2/receipts'
screenshots_dir = '/Users/luna-openclaw/Desktop/spesenabrechnungen/email2/screenshots'
pdf_imgs_dir = '/Users/luna-openclaw/Desktop/spesenabrechnungen/email2/pdf_images'
os.makedirs(pdf_imgs_dir, exist_ok=True)

def pdf_to_pngs(pdf_path, prefix):
    doc = fitz.open(pdf_path)
    paths = []
    for i, page in enumerate(doc):
        mat = fitz.Matrix(1.5, 1.5)
        pix = page.get_pixmap(matrix=mat)
        out = os.path.join(pdf_imgs_dir, f'{prefix}_p{i+1}.png')
        pix.save(out)
        paths.append(out)
    doc.close()
    return paths

def ss(name):
    p = os.path.join(screenshots_dir, f'{name}.png')
    return p if os.path.exists(p) else None

def receipt_pdf(partial):
    for f in os.listdir(receipts_dir):
        if partial in f and f.endswith('.pdf'):
            return os.path.join(receipts_dir, f)
    return None

lines = [
    (1,  '30.01.2026', '116,20 EUR', 'FREENOW DEUTSCHLAND',          ('placeholder', 'Freenow-Quittung 30.01.2026 (116,20 EUR) nicht in Belegen enthalten')),
    (2,  '02.02.2026', '20,00 EUR',  'FREE NOW',                      ('screenshot', ss('070_2026-02-03_0002'))),
    (3,  '02.02.2026', '74,98 EUR',  'WOLT (The Klub Kitchen Pestalozzistr.)', ('pdf', receipt_pdf('The_Klub_Kitchen_Pestalozzistr_2026-02-02'))),
    (4,  '03.02.2026', '17,80 EUR',  'FREE NOW',                      ('screenshot', ss('068_2026-02-03_2339'))),
    (5,  '03.02.2026', '56,46 EUR',  'WOLT (Gym Cook Gärtnerplatz)',('pdf', receipt_pdf('Gym_Cook_-_Gärtnerplatz_2026-02-03'))),
    (6,  '04.02.2026', '19,50 EUR',  'FREE NOW',                      ('screenshot', ss('066_2026-02-05_0105'))),
    (7,  '04.02.2026', '68,90 EUR',  'WOLT (Green Club Haidhausen)',  ('pdf', receipt_pdf('Green_Club_Haidhausen_2026-02-04'))),
    (8,  '10.02.2026', '18,00 EUR',  'FREENOW DEUTSCHLAND',           ('screenshot', ss('064_2026-02-10_0033'))),
    (9,  '09.02.2026', '94,98 EUR',  'WOLT (The Klub Kitchen Pestalozzistr.)', ('pdf', receipt_pdf('The_Klub_Kitchen_Pestalozzistr_2026-02-09'))),
    (10, '10.02.2026', '1,50 EUR',   'FREE NOW',                      ('placeholder', 'Freenow-Quittung 10.02.2026 (1,50 EUR) nicht in Belegen enthalten')),
    (11, '10.02.2026', '85,00 EUR',  'WOLT (Gym Cook Gärtnerplatz)',('pdf', receipt_pdf('Gym_Cook_-_Gärtnerplatz_2026-02-10'))),
    (12, '12.02.2026', '3,58 EUR',   'FREENOW DEUTSCHLAND',           ('placeholder', 'Freenow-Quittung 12.02.2026 (3,58 EUR) nicht in Belegen enthalten')),
    (13, '15.02.2026', '13,60 EUR',  'FREE NOW',                      ('screenshot', ss('062_2026-02-15_2130'))),
    (14, '16.02.2026', '17,80 EUR',  'FREE NOW',                      ('screenshot', ss('060_2026-02-16_2342'))),
    (15, '16.02.2026', '112,08 EUR', 'WOLT (Pepenero Glockenbach)',   ('pdf', receipt_pdf('Pepenero_Glockenbach_2026-02-16'))),
    (16, '17.02.2026', '19,20 EUR',  'FREE NOW',                      ('screenshot', ss('059_2026-02-17_2308'))),
    (17, '23.02.2026', '106,55 EUR', 'DEAN&DAVID F T185 GMBH',       ('screenshot', ss('052_2026-02-23_0901'))),
    (18, '23.02.2026', '13,90 EUR',  'FREE NOW',                      ('screenshot', ss('049_2026-02-23_1958'))),
    (19, '24.02.2026', '124,95 EUR', 'MUNICH CHAUFFEURS GMBH',        ('screenshot', ss('047_2026-02-24_0924'))),
    (20, '23.02.2026', '160,00 EUR', 'RISTORANTE ISOLETTA',           ('img', os.path.join(receipts_dir, '2026-02-23_1757_IMG_0210.jpg'))),
    (21, '24.02.2026', '94,81 EUR',  'WOLT (Stadtsalat Frankfurt)',   ('pdf', receipt_pdf('Stadtsalat_Frankfurt_2026-02-24'))),
    (22, '25.02.2026', '4,76 EUR',   'BCD TRAVEL (Beleg 29939839)',   ('pdf', receipt_pdf('de1_2026_0000_29939839'))),
    (23, '25.02.2026', '4,76 EUR',   'BCD TRAVEL (Beleg 29939837)',   ('pdf', receipt_pdf('de1_2026_0000_29939837'))),
    (24, '24.02.2026', '102,89 EUR', 'DEUTSCHE BAHN Kaufbeleg 925552365186', ('pdf', receipt_pdf('DB_Kaufbeleg_925552365186'))),
    (25, '25.02.2026', '44,70 EUR',  'WOLT (District Banh Mi)',       ('pdf', receipt_pdf('District_Banh_Mi_-_Vietnamese_Streetfood_2026-02-25'))),
    (26, '24.02.2026', '313,82 EUR', 'DEUTSCHE LUFTHANSA BSP',        ('pdf', receipt_pdf('Receipt_1110-971-544'))),
    (27, '24.02.2026', '366,63 EUR', 'DEUTSCHE LUFTHANSA BSP (2)',    ('pdf', receipt_pdf('download'))),
    (28, '26.02.2026', '13,20 EUR',  'FREE NOW',                      ('screenshot', ss('041_2026-02-27_0547'))),
    (29, '27.02.2026', '15,00 EUR',  'ES GLOBAL EXIM GMBH (Taxi FFM)',('img', os.path.join(receipts_dir, '2026-03-03_2011_IMG_0216.jpeg'))),
    (30, '27.02.2026', '18,50 EUR',  'FREE NOW',                      ('screenshot', ss('039_2026-02-27_2140'))),
    (31, '27.02.2026', '24,54 EUR',  'WOLT (Gym Cook Gärtnerplatz)',('pdf', receipt_pdf('Gym_Cook_-_Gärtnerplatz_2026-02-27'))),
    (32, '27.02.2026', '732,00 EUR', 'MARRIOTT HOTEL FRANKFURT',      ('pdf', receipt_pdf('berwi_folio_g111582522'))),
    (33, '03.03.2026', '2,80 CHF / 3,17 EUR', 'VBZ AUTOMATEN (Zürich Tram)', ('placeholder', 'VBZ Automaten Quittung (Zürich Tram, 2,80 CHF) nicht in Belegen enthalten')),
    (34, '02.03.2026', '136,85 EUR', 'MUNICH CHAUFFEURS GMBH',        ('screenshot', ss('038_2026-03-02_1855'))),
    (35, '02.03.2026', '482,25 CHF / 545,81 EUR', 'STRIPE PAYMENTS (Hotel Zürich)', ('screenshot', ss('022_2026-03-17_2200'))),
    (36, '03.03.2026', '2,80 CHF / 3,16 EUR', 'VBZ AUTOMATEN (#2)',   ('placeholder', 'VBZ Automaten Quittung #2 (Zürich Tram, 2,80 CHF) nicht in Belegen enthalten')),
    (37, '05.03.2026', '12,00 CHF / 13,58 EUR', 'UBER',               ('screenshot', ss('033_2026-03-05_0304'))),
    (38, '05.03.2026', '13,50 CHF / 15,28 EUR', 'UBER',               ('screenshot', ss('032_2026-03-05_0555'))),
    (39, '05.03.2026', '2,80 CHF / 3,17 EUR', 'VBZ AUTOMATEN (#3)',   ('placeholder', 'VBZ Automaten Quittung #3 (Zürich Tram, 2,80 CHF) nicht in Belegen enthalten')),
    (40, '05.03.2026', '42,26 CHF / 47,84 EUR', 'UBER',               ('screenshot', ss('031_2026-03-05_1635'))),
    (41, '05.03.2026', '124,95 EUR', 'MUNICH CHAUFFEURS GMBH',        ('screenshot', ss('030_2026-03-05_2008'))),
    (42, '10.03.2026', '17,80 EUR',  'FREE NOW',                      ('screenshot', ss('028_2026-03-10_2249'))),
    (43, '10.03.2026', '20,50 EUR',  'WOLT (Gym Cook Gärtnerplatz)',('pdf', receipt_pdf('Gym_Cook_-_Gärtnerplatz_2026-03-10'))),
    (44, '11.03.2026', '29,72 EUR',  'REWE 0591',                     ('placeholder', 'REWE-Quittung 11.03.2026 (29,72 EUR) nicht in Belegen enthalten')),
    (45, '11.03.2026', '39,00 EUR',  'WOLT (The Klub Kitchen Pestalozzistr.)', ('pdf', receipt_pdf('The_Klub_Kitchen_Pestalozzistr_2026-03-11'))),
    (46, '12.03.2026', '16,80 EUR',  'FREE NOW',                      ('screenshot', ss('026_2026-03-12_2148'))),
    (47, '16.03.2026', '17,80 EUR',  'FREE NOW',                      ('screenshot', ss('024_2026-03-16_2325'))),
    (48, '16.03.2026', '111,38 EUR', 'WOLT (Viva Maria)',             ('pdf', receipt_pdf('Viva_Maria_2026-03-16'))),
    (49, '17.03.2026', '17,60 EUR',  'FREENOW DEUTSCHLAND',           ('screenshot', ss('021_2026-03-17_2323'))),
    (50, '17.03.2026', '84,83 EUR',  'WOLT (The Klub Kitchen Pestalozzistr.)', ('pdf', receipt_pdf('The_Klub_Kitchen_Pestalozzistr_2026-03-17'))),
    (51, '18.03.2026', '20,20 EUR',  'FREE NOW',                      ('screenshot', ss('020_2026-03-18_1926'))),
    (52, '18.03.2026', '30,80 EUR',  'FREE NOW',                      ('screenshot', ss('019_2026-03-18_2051'))),
    (53, '18.03.2026', '40,20 EUR',  'FREE NOW',                      ('screenshot', ss('018_2026-03-18_2330'))),
    (54, '24.03.2026', '4,76 EUR',   'BCD TRAVEL (Beleg 30146801)',   ('pdf', receipt_pdf('de1_2026_0000_30146801'))),
    (55, '23.03.2026', '40,19 EUR',  'WOLT (The Klub Kitchen Pestalozzistr.)', ('pdf', receipt_pdf('The_Klub_Kitchen_Pestalozzistr_2026-03-23'))),
    (56, '24.03.2026', '14,90 EUR',  'FREE NOW',                      ('screenshot', ss('013_2026-03-25_0025'))),
    (57, '24.03.2026', '86,78 EUR',  'WOLT (Gym Cook Gärtnerplatz)',('pdf', receipt_pdf('Gym_Cook_-_Gärtnerplatz_2026-03-24'))),
    (58, '23.03.2026', '313,08 EUR', 'DEUTSCHE LUFTHANSA BSP',        ('screenshot', ss('017_2026-03-23_1440'))),
    (59, '25.03.2026', '17,00 EUR',  'FREE NOW',                      ('screenshot', ss('010_2026-03-26_2142'))),
    (60, '25.03.2026', '95,78 EUR',  "WOLT (Ruff's Burger Rindermarkt)", ('pdf', receipt_pdf("Ruff's_Burger_Rindermarkt_2026-03-25"))),
    (61, '26.03.2026', '18,60 EUR',  'FREE NOW',                      ('screenshot', ss('011_2026-03-26_0141'))),
    (62, '29.03.2026', '70,00 EUR',  'FIRMA DR. PRADHAN (Taxi)',       ('placeholder', 'Taxi-Quittung Dr. Pradhan 29.03.2026 (70,00 EUR) nicht in Belegen enthalten')),
    (63, '29.03.2026', '124,95 EUR', 'MUNICH CHAUFFEURS GMBH',        ('screenshot', ss('009_2026-03-29_1538'))),
    (64, '30.03.2026', '13,20 EUR',  'FREE NOW',                      ('screenshot', ss('006_2026-03-30_1852'))),
    (65, '30.03.2026', '22,00 EUR',  'FREE NOW',                      ('screenshot', ss('005_2026-03-30_2045'))),
    (66, '30.03.2026', '22,60 EUR',  'FREE NOW',                      ('screenshot', ss('004_2026-03-31_0527'))),
    (67, '30.03.2026', '218,89 EUR', 'DEUTSCHE BAHN Ticket 376577215246', ('pdf', receipt_pdf('Ticket_376577215246'))),
    (68, '30.03.2026', '-218,89 EUR','DEUTSCHE BAHN Kaufbeleg/Storno 376577215246', ('pdf', receipt_pdf('DB_Kaufbeleg_376577215246'))),
    (69, '30.03.2026', '236,89 EUR', 'DEUTSCHE BAHN neues Ticket',    ('placeholder', 'Neues DB Ticket 236,89 EUR nach Storno - kein separater Beleg vorhanden; siehe Pos. 67/68')),
]

# Build document
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

title = doc.add_heading('Spesenabrechnung Simon Baumann – SWLF-0020 (Billable)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_width = Inches(6.5)
img_width = Inches(6.0)

for line_num, date, amount, merchant, source in lines:
    src_type, src_val = source

    # Section heading
    heading_text = f'Position {line_num} | {date} | {amount} | {merchant}'
    h = doc.add_heading(heading_text, level=1)

    if src_type == 'placeholder':
        p = doc.add_paragraph()
        run = p.add_run(f'⚠️ PLATZHALTER: {src_val}')
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        print(f'  [PLACEHOLDER] Line {line_num}: {merchant}')

    elif src_type == 'screenshot':
        if src_val and os.path.exists(src_val):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(src_val, width=img_width)
            print(f'  [SS] Line {line_num}: {merchant}')
        else:
            p = doc.add_paragraph()
            run = p.add_run(f'⚠️ PLATZHALTER: Screenshot nicht gefunden ({src_val})')
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    elif src_type == 'img':
        if src_val and os.path.exists(src_val):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(src_val, width=img_width)
            print(f'  [IMG] Line {line_num}: {merchant}')
        else:
            p = doc.add_paragraph()
            run = p.add_run(f'⚠️ PLATZHALTER: Bild nicht gefunden ({src_val})')
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    elif src_type == 'pdf':
        if src_val and os.path.exists(src_val):
            pngs = pdf_to_pngs(src_val, f'line{line_num:02d}')
            for png in pngs:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(png, width=img_width)
            print(f'  [PDF {len(pngs)}p] Line {line_num}: {merchant}')
        else:
            p = doc.add_paragraph()
            run = p.add_run(f'⚠️ PLATZHALTER: PDF nicht gefunden ({src_val})')
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    doc.add_page_break()

out_path = '/Users/luna-openclaw/Desktop/spesenabrechnungen/Spesenabrechnungen_Simon_SWLF0020.docx'
doc.save(out_path)
size_mb = os.path.getsize(out_path) / 1024 / 1024
print(f'\nSaved: {out_path} ({size_mb:.1f} MB)')
